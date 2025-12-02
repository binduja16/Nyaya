import json
import re
import logging
from io import BytesIO
from typing import List, Dict, Any, Optional
from fastapi import FastAPI, UploadFile, Request, Form, Query, File
from fastapi.responses import JSONResponse, HTMLResponse, Response, FileResponse
from fastapi.templating import Jinja2Templates
from fastapi.staticfiles import StaticFiles
from fastapi.middleware.cors import CORSMiddleware
from PyPDF2 import PdfReader
from docx import Document
import httpx
from pydantic import BaseModel
import random
import pandas as pd
import os
from googletrans import Translator
import gtts
from datetime import datetime

# ------------------ Logging Setup (FIRST) ------------------
logging.basicConfig(level=logging.INFO, format="%(asctime)s [%(levelname)s] %(message)s")
logger = logging.getLogger(__name__)

# ------------------ Import Chat Module with Proper Error Handling ------------------
try:
    from chat import process_chat_query, generate_case_duration_response, generate_audio_evidence_response, ChatMessage as ChatModuleMessage, chat_service
    logger.info("Successfully imported chat module")
except ImportError as e:
    logger.error(f"Failed to import chat module: {e}")
    # Define comprehensive fallback functions
    class ChatResponse:
        def __init__(self, reply, language, source="fallback"):
            self.reply = reply
            self.language = language
            self.source = source
    
    async def process_chat_query(query, history=None, language="en"):
        return ChatResponse(
            reply="🚨 Chat module temporarily unavailable. Please try again later.",
            language=language,
            source="fallback"
        )
    
    def generate_case_duration_response():
        return "🚨 Chat module temporarily unavailable. Case duration information not accessible."
    
    def generate_audio_evidence_response(detailed=False):
        return "🚨 Chat module temporarily unavailable. Audio evidence information not accessible."
    
    class ChatModuleMessage:
        def __init__(self, role, content):
            self.role = role
            self.content = content
    
    chat_service = type('obj', (object,), {
        '_generate_fallback_response': lambda self, query: "🚨 Chat service unavailable. Please try again later."
    })()

# ------------------ CONFIG ------------------
# Separate API keys for different services
FIR_ANALYSIS_API_KEY = "AIzaSyCo3AF8OzpCwbdh3Zirhp1DhlCUJwWApeQ"
CHAT_API_KEY = "AIzaSyCqPHhgSPIs3ZEpbVaEQ3485Aqop1tQBZQ"  # Different key for chat
MODEL_URL = "https://generativelanguage.googleapis.com/v1beta/models/gemini-2.5-flash:generateContent"

# ------------------ FastAPI setup ------------------
app = FastAPI(title="AI Courtroom Assistant")
templates = Jinja2Templates(directory="templates")
app.mount("/static", StaticFiles(directory="static"), name="static")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# ------------------ Global Variables ------------------
# Initialize conversation_history as a global variable
conversation_history = []

# ------------------ Case Context ------------------
case_context = {
    "summary": "",
    "draft": "",
    "ipc_sections": [],
    "actions": [],
    "weak_spots": [],
    "advocates": [],
    "fir_number": "",
    "clarification_questions": [],
    "user_location": {"lat": 12.9716, "lng": 77.5946},
    "detected_location": "Unknown"
}

class ChatMessage(BaseModel):
    role: str
    content: str

class ChatRequest(BaseModel):
    query: str

# ------------------ Translation Service ------------------
class TranslationService:
    def __init__(self):
        self.translator = Translator()
        self.supported_languages = {
            'en': 'English',
            'hi': 'Hindi',
            'bn': 'Bengali', 
            'te': 'Telugu',
            'mr': 'Marathi',
            'ta': 'Tamil',
            'ur': 'Urdu',
            'gu': 'Gujarati',
            'kn': 'Kannada',
            'ml': 'Malayalam',
            'pa': 'Punjabi'
        }
    
    async def translate_text(self, text: str, target_lang: str) -> str:
        """Translate text to target language"""
        try:
            if target_lang == 'en' or not text.strip():
                return text
            
            translation = self.translator.translate(text, dest=target_lang)
            return translation.text
        except Exception as e:
            logger.error(f"Translation error: {e}")
            return text
    
    async def translate_legal_content(self, content: Dict, target_lang: str) -> Dict:
        """Translate all legal content to target language"""
        try:
            translated_content = {}
            
            # Translate summary
            if content.get('summary'):
                translated_content['summary'] = await self.translate_text(
                    content['summary'], target_lang
                )
            
            # Translate draft petition
            if content.get('draft'):
                translated_content['draft'] = await self.translate_text(
                    content['draft'], target_lang
                )
            
            # Translate IPC sections
            if content.get('ipc_sections'):
                translated_sections = []
                for section in content['ipc_sections']:
                    if isinstance(section, str):
                        translated_sections.append(
                            await self.translate_text(section, target_lang)
                        )
                    elif isinstance(section, dict):
                        translated_section = {}
                        for key, value in section.items():
                            translated_section[key] = await self.translate_text(
                                str(value), target_lang
                            ) if isinstance(value, str) else value
                        translated_sections.append(translated_section)
                translated_content['ipc_sections'] = translated_sections
            
            # Translate actions
            if content.get('actions'):
                translated_actions = []
                for action in content['actions']:
                    translated_action = {}
                    for key, value in action.items():
                        translated_action[key] = await self.translate_text(
                            str(value), target_lang
                        ) if isinstance(value, str) else value
                    translated_actions.append(translated_action)
                translated_content['actions'] = translated_actions
            
            # Translate weak spots
            if content.get('weak_spots'):
                translated_weak_spots = []
                for weak_spot in content['weak_spots']:
                    translated_weak_spot = {}
                    for key, value in weak_spot.items():
                        translated_weak_spot[key] = await self.translate_text(
                            str(value), target_lang
                        ) if isinstance(value, str) else value
                    translated_weak_spots.append(translated_weak_spot)
                translated_content['weak_spots'] = translated_weak_spots
            
            # Translate advocates info
            if content.get('advocates'):
                translated_advocates = []
                for advocate in content['advocates']:
                    translated_advocate = advocate.copy()
                    # Translate text fields
                    text_fields = ['name', 'address', 'experience', 'location', 'price_range']
                    for field in text_fields:
                        if field in translated_advocate and translated_advocate[field]:
                            translated_advocate[field] = await self.translate_text(
                                translated_advocate[field], target_lang
                            )
                    
                    # Translate specialities
                    if 'speciality' in translated_advocate and isinstance(translated_advocate['speciality'], list):
                        translated_specialities = []
                        for speciality in translated_advocate['speciality']:
                            translated_specialities.append(
                                await self.translate_text(speciality, target_lang)
                            )
                        translated_advocate['speciality'] = translated_specialities
                    
                    translated_advocates.append(translated_advocate)
                translated_content['advocates'] = translated_advocates
            
            return translated_content
            
        except Exception as e:
            logger.error(f"Legal content translation error: {e}")
            return content

# Initialize translation service
translation_service = TranslationService()

# ------------------ Text-to-Speech Service ------------------
class TextToSpeechService:
    def __init__(self):
        self.supported_languages = {
            'en': 'en',
            'hi': 'hi',
            'bn': 'bn',
            'te': 'te', 
            'mr': 'mr',
            'ta': 'ta',
            'ur': 'ur',
            'gu': 'gu',
            'kn': 'kn',
            'ml': 'ml',
            'pa': 'pa'
        }
    
    async def generate_speech(self, text: str, language: str) -> bytes:
        """Generate speech audio from text"""
        try:
            if language not in self.supported_languages:
                language = 'en'
            
            # Clean text for TTS
            clean_text = re.sub(r'[^\w\s.,!?;:]', '', text)
            if len(clean_text) > 500:  # Limit text length for TTS
                clean_text = clean_text[:500] + "..."
            
            tts = gtts.gTTS(text=clean_text, lang=language, slow=False)
            audio_buffer = BytesIO()
            tts.write_to_fp(audio_buffer)
            audio_buffer.seek(0)
            return audio_buffer.getvalue()
            
        except Exception as e:
            logger.error(f"TTS generation error: {e}")
            return b''

tts_service = TextToSpeechService()

# ------------------ Indian Cities Mapping ------------------
INDIAN_CITIES = {
    "delhi": {"lat": 28.6139, "lng": 77.2090},
    "mumbai": {"lat": 19.0760, "lng": 72.8777},
    "bangalore": {"lat": 12.9716, "lng": 77.5946},
    "chennai": {"lat": 13.0827, "lng": 80.2707},
    "kolkata": {"lat": 22.5726, "lng": 88.3639},
    "hyderabad": {"lat": 17.3850, "lng": 78.4867},
    "pune": {"lat": 18.5204, "lng": 73.8567},
    "ahmedabad": {"lat": 23.0225, "lng": 72.5714},
    "jaipur": {"lat": 26.9124, "lng": 75.7873},
    "lucknow": {"lat": 26.8467, "lng": 80.9462},
    "bengaluru": {"lat": 12.9716, "lng": 77.5946}
}

# ------------------ Load Advocates from CSV ------------------
def load_advocate_csv() -> List[Dict[str, Any]]:
    """Load advocate data from CSV file with proper column mapping"""
    try:
        csv_file = "advocate-UK-almora-aug-16.csv"
        logger.info(f"Looking for CSV file: {csv_file}")
        
        if not os.path.exists(csv_file):
            logger.warning(f"CSV file not found: {csv_file}, using sample data")
            return get_sample_advocates()
        
        # Read CSV with proper column names
        df = pd.read_csv(csv_file, header=None, names=[
            'City', 'CaseNumber1', 'CaseNumber2', 'Name', 'Address', 'Date', 'Status'
        ])
        
        logger.info(f"Successfully loaded CSV with {len(df)} rows")
        
        # Convert DataFrame to list of dictionaries with enhanced information
        advocates = []
        for index, row in df.iterrows():
            try:
                # Extract city from the data
                city = str(row['City']).strip().title() if pd.notna(row['City']) else "Unknown"
                
                # Create proper advocate profile
                advocate = {
                    "name": f"Adv. {str(row['Name']).strip()}" if pd.notna(row['Name']) else f"Advocate {index+1}",
                    "enrollment_number": f"EN{index+1:04d}",
                    "address": str(row['Address']).strip() if pd.notna(row['Address']) else "Address not specified",
                    "phone": generate_phone_number(),
                    "email": generate_email(str(row['Name']).strip() if pd.notna(row['Name']) else f"advocate{index+1}"),
                    "speciality": ["Criminal Law", "Civil Law", "General Practice"],
                    "experience": f"{random.randint(3, 20)} years",
                    "location": city,
                    "rating": round(3.5 + random.random() * 1.5, 1),
                    "price_range": generate_price_range(),
                    "case_number": str(row['CaseNumber1']).strip() if pd.notna(row['CaseNumber1']) else "N/A"
                }
                advocates.append(advocate)
                
            except Exception as row_error:
                logger.error(f"Error processing row {index}: {row_error}")
                continue
        
        logger.info(f"Successfully processed {len(advocates)} advocates")
        return advocates
        
    except Exception as e:
        logger.error(f"Error loading advocate CSV: {e}")
        return get_sample_advocates()

def get_sample_advocates():
    """Generate sample advocates if CSV is not available"""
    sample_advocates = []
    cities = ["Delhi", "Mumbai", "Bangalore", "Chennai", "Kolkata", "Hyderabad", "Pune"]
    
    for i in range(8):
        advocate = {
            "name": f"Adv. Sample {i+1}",
            "enrollment_number": f"EN{i+1:04d}",
            "address": f"Sample Address {i+1}, {random.choice(cities)}",
            "phone": generate_phone_number(),
            "email": f"advocate{i+1}@example.com",
            "speciality": ["Criminal Law", "Civil Law"],
            "experience": f"{random.randint(5, 25)} years",
            "location": random.choice(cities),
            "rating": round(3.5 + random.random() * 1.5, 1),
            "price_range": generate_price_range(),
            "case_number": f"CASE{random.randint(1000, 9999)}"
        }
        sample_advocates.append(advocate)
    
    return sample_advocates

def generate_phone_number():
    """Generate a realistic Indian phone number"""
    prefixes = ["98", "99", "97", "96", "95", "94", "93", "92", "91", "90"]
    return f"+91-{random.choice(prefixes)}{random.randint(10000000, 99999999)}"

def generate_email(name):
    """Generate email from name"""
    clean_name = re.sub(r'[^a-zA-Z]', '', name).lower()
    domains = ["law.com", "legal.com", "advocate.com", "gmail.com"]
    return f"{clean_name}@{random.choice(domains)}"

def generate_price_range():
    """Generate realistic price range for advocates"""
    ranges = [
        "₹2,000 - ₹8,000",
        "₹3,000 - ₹10,000", 
        "₹5,000 - ₹15,000",
        "₹8,000 - ₹20,000",
        "₹10,000 - ₹25,000"
    ]
    return random.choice(ranges)

# ------------------ Helper Functions ------------------
def extract_text(file_bytes: BytesIO) -> str:
    """Extract text from PDF, DOCX, or TXT"""
    try:
        reader = PdfReader(file_bytes)
        text = "\n".join([page.extract_text() or "" for page in reader.pages])
        if text.strip():
            return text
    except Exception as e:
        logger.warning(f"PDF extraction failed: {e}")
    
    try:
        doc = Document(file_bytes)
        text = "\n".join([p.text for p in doc.paragraphs])
        if text.strip():
            return text
    except Exception as e:
        logger.warning(f"DOCX extraction failed: {e}")
    
    try:
        file_bytes.seek(0)
        return file_bytes.read().decode("utf-8", errors="ignore")
    except Exception as e:
        logger.error(f"Text extraction failed: {e}")
        return ""

def extract_json_from_text(text: str):
    """Extract and parse JSON safely from Gemini output"""
    if not text:
        return None
        
    text = re.sub(r"```(?:json)?\n?", "", text)
    text = re.sub(r"```", "", text)
    try:
        return json.loads(text)
    except Exception:
        match = re.search(r"\{[\s\S]*\}", text)
        if match:
            try:
                return json.loads(match.group(0))
            except Exception:
                pass
    return None

async def analyze_fir_with_gemini(fir_text: str):
    """Analyze FIR text using Gemini and extract legal insights."""
    if not fir_text or len(fir_text.strip()) < 50:
        return {
            "ipc_sections": [],
            "draft": "FIR text is too short to analyze.",
            "actions": [],
            "weak_spots": []
        }

    fir_text = fir_text[:4000]
    prompt = f"""
You are an expert Indian legal assistant and FIR analyzer.

Read the FIR text below and provide a **structured legal summary** in **pure JSON** format only.

Your output must strictly follow this structure:
{{
  "ipc_sections": ["List relevant IPC sections (e.g., 498A, 506, etc.)"],
  "draft": "Draft legal petition or complaint in formal language (100-200 words)",
  "actions": [
    {{
      "step": "Action Title",
      "description": "What the complainant or advocate should do next"
    }}
  ],
  "weak_spots": [
    {{
      "type": "Weakness Category",
      "description": "Detailed description of the weakness"
    }}
  ]
}}

Analyze thoroughly and ensure IPC sections are correct under Indian law.

FIR TEXT:
{fir_text}
"""

    headers = {"x-goog-api-key": FIR_ANALYSIS_API_KEY, "Content-Type": "application/json"}
    data = {
        "contents": [{"parts": [{"text": prompt}]}],
        "generationConfig": {"temperature": 0.3, "maxOutputTokens": 3500}
    }

    try:
        async with httpx.AsyncClient(timeout=150) as client:
            resp = await client.post(MODEL_URL, headers=headers, json=data)
            resp.raise_for_status()
            content = resp.json()
            text = (
                content.get("candidates", [{}])[0]
                .get("content", {}).get("parts", [{}])[0].get("text", "")
            )
    except Exception as e:
        logger.error(f"Gemini API error for FIR analysis: {e}")
        return {
            "ipc_sections": [], 
            "draft": "Error contacting Gemini for FIR analysis.", 
            "actions": [],
            "weak_spots": []
        }

    json_data = extract_json_from_text(text)
    if not json_data:
        logger.warning("Invalid or empty Gemini output for FIR analysis.")
        return {
            "ipc_sections": [], 
            "draft": "Unable to generate draft from FIR analysis.", 
            "actions": [],
            "weak_spots": []
        }

    return json_data

def detect_weak_spots(text: str) -> List[Dict[str, str]]:
    """Enhanced weak spot detection with structured output"""
    weak_points = []
    
    if not re.search(r"\b\d{1,2}[-/]\d{1,2}[-/]\d{2,4}\b", text):
        weak_points.append({
            "type": "Missing Date", 
            "description": "No clear incident date mentioned. This is crucial for establishing timeline."
        })
    
    location_indicators = ["road", "street", "house", "building", "area", "sector", "colony"]
    if not any(indicator in text.lower() for indicator in location_indicators):
        weak_points.append({
            "type": "Vague Location", 
            "description": "Incident location details are insufficient. Add specific address, landmarks, or PIN code."
        })
    
    if not re.search(r"(hit|beat|slap|punch|kick|push|threaten|abuse|harass|force)", text, re.I):
        weak_points.append({
            "type": "Lack of Specifics", 
            "description": "Details of the actual incident are unclear. Describe exactly what happened."
        })
    
    evidence_indicators = ["witness", "evidence", "proof", "photo", "video", "document"]
    if not any(indicator in text.lower() for indicator in evidence_indicators):
        weak_points.append({
            "type": "Evidence Gap", 
            "description": "No mention of supporting evidence, witnesses, or documentation."
        })
    
    return weak_points

def generate_dynamic_questions(fir_text: str) -> List[Dict]:
    """Generate clarification questions based on FIR content"""
    questions = []
    
    if not re.search(r"\b\d{1,2}[-/]\d{1,2}[-/]\d{2,4}\b", fir_text):
        questions.append({
            "question": "On what exact date did the incident occur? (DD/MM/YYYY format)",
            "field": "date"
        })
    
    location_indicators = ["road", "street", "house", "building", "area"]
    if not any(indicator in fir_text.lower() for indicator in location_indicators):
        questions.append({
            "question": "Where exactly did the incident take place? Please provide full address with landmarks.",
            "field": "location"
        })
    
    return questions[:3]

async def detect_location_from_text(text: str) -> Dict[str, float]:
    """Detect location from FIR text using city names"""
    text_lower = text.lower()
    
    for city, coords in INDIAN_CITIES.items():
        if city in text_lower:
            logger.info(f"Detected location: {city}")
            return coords
    
    location_keywords = {
        "delhi": ["delhi", "new delhi", "ncr"],
        "mumbai": ["mumbai", "bombay"],
        "bangalore": ["bangalore", "bengaluru", "blr"],
        "chennai": ["chennai", "madras"],
        "kolkata": ["kolkata", "calcutta"],
        "hyderabad": ["hyderabad"],
        "pune": ["pune"]
    }
    
    for city, keywords in location_keywords.items():
        if any(keyword in text_lower for keyword in keywords):
            logger.info(f"Detected location via keywords: {city}")
            return INDIAN_CITIES[city]
    
    return INDIAN_CITIES["bengaluru"]

async def get_nearby_advocates(lat: float, lng: float, ipc_sections: List[str]) -> List[Dict]:
    """Get advocates from CSV based on location and case type"""
    try:
        all_advocates = load_advocate_csv()
        if not all_advocates:
            return []
        
        speciality = determine_speciality(ipc_sections)
        city = get_city_from_coordinates(lat, lng)
        logger.info(f"Searching advocates in: {city}")
        
        filtered_advocates = []
        for advocate in all_advocates:
            if city.lower() in advocate["location"].lower():
                enhanced_advocate = advocate.copy()
                enhanced_advocate["speciality"] = speciality
                filtered_advocates.append(enhanced_advocate)
        
        if not filtered_advocates:
            filtered_advocates = all_advocates[:8]
        
        for advocate in filtered_advocates:
            advocate["speciality"] = speciality
        
        return filtered_advocates[:8]
        
    except Exception as e:
        logger.error(f"Error fetching advocates from CSV: {e}")
        return []

def determine_speciality(ipc_sections: List[str]) -> List[str]:
    """Determine advocate speciality based on IPC sections"""
    specialities = []
    
    ipc_to_speciality = {
        "420": ["Fraud Cases", "Criminal Law"],
        "506": ["Criminal Intimidation", "Criminal Law"],
        "498A": ["Domestic Violence", "Women Rights"],
        "323": ["Assault", "Criminal Law"],
        "354": ["Sexual Harassment", "Women Rights"],
        "376": ["Rape", "Criminal Law"],
        "406": ["Criminal Breach of Trust", "Criminal Law"]
    }
    
    for section in ipc_sections:
        section_num = re.search(r'\b(\d+[A-Z]*)\b', section)
        if section_num:
            section_key = section_num.group(1)
            if section_key in ipc_to_speciality:
                specialities.extend(ipc_to_speciality[section_key])
    
    specialities = list(set(specialities))
    if not specialities:
        specialities = ["Criminal Law", "General Practice"]
    
    if "Criminal Law" not in specialities:
        specialities.append("Criminal Law")
    
    return specialities[:3]

def get_city_from_coordinates(lat: float, lng: float) -> str:
    """Get city name from coordinates"""
    for city, coords in INDIAN_CITIES.items():
        if (abs(coords["lat"] - lat) < 1.0 and abs(coords["lng"] - lng) < 1.0):
            return city.capitalize()
    return "Bengaluru"

# ------------------ Complaint Generation Functions ------------------
def build_prompt(data: dict) -> str:
    date_str = data.get("incident_date") or datetime.now().strftime("%d %B %Y")

    prompt = f"""
DRAFT POLICE COMPLAINT (INDIA)

You are a professional legal drafting assistant. Prepare a formal written complaint suitable to submit to a Police Station in India.

Use **only** the user's data — do not invent facts. Follow Indian police complaint format:
- Heading: "To The Officer In-Charge, [Police Station Name]"
- Parties: Complainant / Accused
- Numbered facts of incident
- Prayer / Relief section
- Annexure / Evidence list
- Verification (Affidavit style)

COMPLAINT DETAILS:
Complainant name: {data.get('petitioner_name','')}
Complainant address: {data.get('petitioner_address','')}
Accused / Respondent: {data.get('respondent_name','')}
Accused address: {data.get('respondent_address','')}
Police Station: {data.get('place','')}
Date of incident: {date_str}
Facts in detail: {data.get('facts','')}
Relief sought: {data.get('relief','')}
Sections mentioned (if any): {data.get('sections','')}
Evidence / documents / witnesses: {data.get('evidence','')}

TASKS:
1. Generate a legally formatted Indian Police Station complaint.
2. Include heading, cause title, numbered facts, prayer, annexure list, and verification.
3. Suggest IPC/CrPC sections **only if relevant**. Do not give legal advice.
4. Output should be plain formatted text — NO MARKDOWN.

End with this disclaimer:
"This is a draft. Consult a practising advocate before filing."

START DRAFT BELOW:
"""
    return prompt

async def call_gemini_generate(prompt_text: str) -> str:
    """Call Gemini API for content generation"""
    try:
        headers = {
            "x-goog-api-key": FIR_ANALYSIS_API_KEY, 
            "Content-Type": "application/json"
        }
        data = {
            "contents": [{"parts": [{"text": prompt_text}]}],
            "generationConfig": {"temperature": 0.3, "maxOutputTokens": 3500}
        }
        
        async with httpx.AsyncClient(timeout=150) as client:
            resp = await client.post(MODEL_URL, headers=headers, json=data)
            resp.raise_for_status()
            content = resp.json()
            text = (
                content.get("candidates", [{}])[0]
                .get("content", {}).get("parts", [{}])[0].get("text", "")
            )
            return text
    except Exception as e:
        logger.error(f"Gemini generation error: {e}")
        return f"[ERROR] Complaint generation failed: {e}"

def create_docx_from_text(title: str, body_text: str) -> BytesIO:
    """Create DOCX file from text"""
    doc = Document()
    doc.add_heading(title, level=1)

    for line in body_text.splitlines():
        if line.strip() == "":
            doc.add_paragraph("")
        else:
            doc.add_paragraph(line)

    file_stream = BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    return file_stream

# ------------------ Routes ------------------
@app.get("/", response_class=HTMLResponse)
async def start(request: Request):
    return templates.TemplateResponse("main.html", {"request": request})

@app.get("/upload", response_class=HTMLResponse)
async def upload_page(request: Request):
    return templates.TemplateResponse("index.html", {"request": request})

@app.post("/process/")
async def process_file(file: UploadFile = File(...), request: Request = None):
    """Process uploaded FIR document and return result.html"""
    if not file.filename.lower().endswith((".pdf", ".docx", ".txt")):
        return JSONResponse({"error": "Only PDF, DOCX, TXT supported"}, status_code=400)

    try:
        file_bytes = await file.read()
        text = extract_text(BytesIO(file_bytes))
        logger.info(f"Extracted text length: {len(text)}")
        
        if not text.strip():
            return JSONResponse({"error": "No text found in file"}, status_code=400)

        logger.info("Starting FIR analysis with Gemini...")
        result = await analyze_fir_with_gemini(text)
        logger.info(f"Gemini analysis result: {result}")
        
        additional_weak_spots = detect_weak_spots(text)
        all_weak_spots = additional_weak_spots + result.get("weak_spots", [])
        
        user_location = await detect_location_from_text(text)
        detected_city = get_city_from_coordinates(user_location["lat"], user_location["lng"])
        
        advocates = await get_nearby_advocates(
            user_location["lat"], 
            user_location["lng"],
            result.get("ipc_sections", [])
        )

        clarification_questions = generate_dynamic_questions(text)

        # Update case context with analysis results
        case_context.update({
            "summary": text[:2000],
            "draft": result.get("draft", ""),
            "ipc_sections": result.get("ipc_sections", []),
            "actions": result.get("actions", []),
            "weak_spots": all_weak_spots,
            "advocates": advocates,
            "clarification_questions": clarification_questions,
            "fir_number": f"FIR{hash(text) % 10000:04d}",
            "user_location": user_location,
            "detected_location": detected_city
        })

        # Return the result.html template with all analysis data
        return templates.TemplateResponse("result.html", {
            "request": request,
            "summary": case_context["summary"],
            "draft": case_context["draft"],
            "ipc_sections": case_context["ipc_sections"],
            "actions": case_context["actions"],
            "weak_spots": case_context["weak_spots"],
            "advocates": case_context["advocates"],
            "clarification_questions": case_context["clarification_questions"],
            "fir_number": case_context["fir_number"],
            "detected_location": case_context["detected_location"]
        })

    except Exception as e:
        logger.error(f"Processing error: {e}")
        return JSONResponse({"error": f"Processing failed: {str(e)}"}, status_code=500)

@app.get("/petition", response_class=HTMLResponse)
async def petition_form(request: Request):
    """Petition form page"""
    return templates.TemplateResponse("petition_form.html", {"request": request})

# Replace your current /generate route with this properly named version
@app.post("/generate", name="generate_complaint")
async def generate_complaint(
    petitioner_name: str = Form(...),
    petitioner_address: str = Form(...),
    respondent_name: str = Form(...),
    respondent_address: str = Form(...),
    place: str = Form(...),
    incident_date: str = Form(""),
    facts: str = Form(...),
    relief: str = Form(...),
    sections: str = Form(""),
    evidence: str = Form("")
):
    """Generate police complaint document"""
    try:
        data = {
            "petitioner_name": petitioner_name,
            "petitioner_address": petitioner_address,
            "respondent_name": respondent_name,
            "respondent_address": respondent_address,
            "place": place,
            "incident_date": incident_date,
            "facts": facts,
            "relief": relief,
            "sections": sections,
            "evidence": evidence
        }
        
        # Build prompt for Gemini
        prompt = build_prompt(data)
        
        # Call Gemini
        generated_text = await call_gemini_generate(prompt)
        
        # Create document
        filename = f"police_complaint_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        file_stream = create_docx_from_text("POLICE COMPLAINT", generated_text)
        
        # Save to temporary storage
        os.makedirs("temp_outputs", exist_ok=True)
        filepath = os.path.join("temp_outputs", filename)
        
        with open(filepath, "wb") as f:
            f.write(file_stream.getvalue())
        
        return JSONResponse({
            "status": "success",
            "draft_text": generated_text,
            "download_filename": filename,
            "message": "Complaint generated successfully"
        })
        
    except Exception as e:
        logger.error(f"Complaint generation error: {e}")
        return JSONResponse(
            {"status": "error", "message": f"Generation failed: {str(e)}"},
            status_code=500
        )

# Also add name to download route
@app.get("/download/{filename}", name="download_file")
async def download_file(filename: str):
    """Download generated document"""
    try:
        file_path = os.path.join("temp_outputs", filename)
        if not os.path.exists(file_path):
            return JSONResponse(
                {"error": "File not found"}, 
                status_code=404
            )
        
        return FileResponse(
            file_path,
            media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            filename=filename
        )
    except Exception as e:
        logger.error(f"Download error: {e}")
        return JSONResponse(
            {"error": "Download failed"}, 
            status_code=500
        )

# Add this test endpoint to debug the routing issue
@app.get("/test-routes")
async def test_routes():
    """Test if all routes are accessible"""
    routes = [
        {"path": "/", "method": "GET", "name": "home"},
        {"path": "/upload", "method": "GET", "name": "upload_page"},
        {"path": "/process/", "method": "POST", "name": "process_file"},
        {"path": "/petition", "method": "GET", "name": "petition_form"},
        {"path": "/generate", "method": "POST", "name": "generate_complaint"},
        {"path": "/download/{filename}", "method": "GET", "name": "download_file"},
        {"path": "/chat/", "method": "POST", "name": "chat_with_ai"}
    ]
    
    return JSONResponse({
        "message": "Available routes",
        "routes": routes,
        "status": "success"
    })

# ------------------ API Endpoints ------------------
@app.post("/chat/")
async def chat_with_ai(request: ChatRequest, language: str = Query('en')):
    """Enhanced chat endpoint using chat.py module"""
    global conversation_history
    
    try:
        user_query = request.query.strip()
        
        if not user_query:
            return {"reply": "Please provide a legal question or query.", "source": "main.py"}
        
        # Convert conversation history to chat.py format
        chat_history = []
        for msg in conversation_history:
            chat_history.append(ChatModuleMessage(role=msg.role, content=msg.content))
        
        # Check for specific queries that have dedicated responses in chat.py
        lower_query = user_query.lower()
        
        # Case duration queries
        duration_keywords = ['how long', 'time taken', 'duration', 'time frame', 'complete case', 'finish case']
        if any(keyword in lower_query for keyword in duration_keywords) and ('case' in lower_query or 'trial' in lower_query):
            response_text = generate_case_duration_response()
        
        # Audio evidence queries
        elif any(keyword in lower_query for keyword in ['audio', 'call recording', 'voice recording']):
            detailed = any(word in lower_query for word in ['explain', 'elaborate', 'detailed', 'more info'])
            response_text = generate_audio_evidence_response(detailed=detailed)
        
        else:
            # Use chat.py service for all other queries
            chat_response = await process_chat_query(user_query, chat_history, language)
            response_text = chat_response.reply
        
        # Update conversation history
        conversation_history.append(ChatMessage(role="user", content=user_query))
        conversation_history.append(ChatMessage(role="assistant", content=response_text))
        
        # Limit conversation history
        if len(conversation_history) > 20:
            conversation_history = conversation_history[-20:]
        
        # Translate if needed (using existing translation service)
        if language != 'en' and response_text:
            try:
                translated_response = await translation_service.translate_text(
                    response_text, language
                )
                return {
                    "reply": translated_response,
                    "original_reply": response_text,
                    "language": language,
                    "source": "chat.py"
                }
            except Exception as translate_error:
                logger.error(f"Translation error: {translate_error}")
                return {
                    "reply": response_text, 
                    "language": "en",
                    "source": "chat.py"
                }
        
        return {
            "reply": response_text, 
            "language": "en",
            "source": "chat.py"
        }
        
    except Exception as e:
        logger.error(f"Chat endpoint error: {e}")
        fallback = chat_service._generate_fallback_response(request.query if request.query else "")
        return {
            "reply": fallback, 
            "language": "en",
            "source": "chat.py"
        }

@app.post("/chat/clear")
async def clear_chat_history():
    """Clear conversation history"""
    global conversation_history
    conversation_history = []
    return {"status": "success", "message": "Chat history cleared"}

@app.get("/chat/history")
async def get_chat_history():
    """Get current conversation history"""
    return {"history": [{"role": msg.role, "content": msg.content} for msg in conversation_history]}

@app.get("/api/case-analysis")
async def get_case_analysis(language: str = Query('en')):
    """API endpoint to get current case analysis with optional translation"""
    base_data = {
        "summary": case_context.get("summary", ""),
        "draft": case_context.get("draft", ""),
        "ipc_sections": case_context.get("ipc_sections", []),
        "actions": case_context.get("actions", []),
        "weak_spots": case_context.get("weak_spots", []),
        "advocates": case_context.get("advocates", []),
        "fir_number": case_context.get("fir_number", ""),
        "detected_location": case_context.get("detected_location", "Unknown")
    }
    
    if language != 'en':
        translated_data = await translation_service.translate_legal_content(
            base_data, language
        )
        return {**base_data, **translated_data, "language": language}
    
    return {**base_data, "language": "en"}

@app.post("/api/translate-content")
async def translate_legal_content(target_lang: str = Form(...)):
    """Translate all legal content to target language"""
    try:
        if target_lang not in translation_service.supported_languages:
            return JSONResponse(
                {"error": f"Unsupported language: {target_lang}"}, 
                status_code=400
            )
        
        case_data = {
            "summary": case_context.get("summary", ""),
            "draft": case_context.get("draft", ""),
            "ipc_sections": case_context.get("ipc_sections", []),
            "actions": case_context.get("actions", []),
            "weak_spots": case_context.get("weak_spots", []),
            "advocates": case_context.get("advocates", []),
        }
        
        translated_data = await translation_service.translate_legal_content(
            case_data, target_lang
        )
        
        return JSONResponse({
            "translated_content": translated_data,
            "target_language": target_lang,
            "language_name": translation_service.supported_languages[target_lang]
        })
        
    except Exception as e:
        logger.error(f"Translation endpoint error: {e}")
        return JSONResponse({"error": "Translation failed"}, status_code=500)

@app.post("/api/text-to-speech")
async def convert_text_to_speech(
    text: str = Form(...),
    language: str = Form('en')
):
    """Convert text to speech in specified language"""
    try:
        audio_data = await tts_service.generate_speech(text, language)
        
        return Response(
            content=audio_data,
            media_type="audio/mpeg",
            headers={"Content-Disposition": "attachment; filename=speech.mp3"}
        )
    except Exception as e:
        logger.error(f"TTS endpoint error: {e}")
        return JSONResponse({"error": "Speech generation failed"}, status_code=500)

@app.get("/api/supported-languages")
async def get_supported_languages():
    """Get list of supported languages"""
    return JSONResponse({
        "languages": translation_service.supported_languages,
        "tts_languages": tts_service.supported_languages
    })

# ------------------ Test Endpoints ------------------
@app.get("/test-advocates")
async def test_advocates():
    """Test endpoint to verify advocate data loading"""
    advocates = load_advocate_csv()
    return JSONResponse({
        "advocates_count": len(advocates),
        "sample_advocate": advocates[0] if advocates else None,
        "cities_available": list(set([adv["location"] for adv in advocates])),
        "status": "success" if advocates else "error"
    })

@app.get("/debug/csv-content")
async def debug_csv_content():
    """Show the actual content of the CSV file"""
    try:
        csv_file = "advocate-UK-almora-aug-16.csv"
        
        if not os.path.exists(csv_file):
            return JSONResponse({"error": f"CSV file not found: {csv_file}"})
        
        with open(csv_file, 'r', encoding='utf-8') as f:
            content = f.read()
        
        try:
            df = pd.read_csv(csv_file, header=None)
            df_info = {
                "rows": len(df),
                "columns": ["City", "CaseNumber1", "CaseNumber2", "Name", "Address", "Date", "Status"],
                "first_3_rows": df.head(3).to_dict('records')
            }
        except Exception as e:
            df_info = {"error": str(e)}
        
        return JSONResponse({
            "file_exists": True,
            "file_size": len(content),
            "raw_content_preview": content[:1000] + "..." if len(content) > 1000 else content,
            "pandas_analysis": df_info
        })
        
    except Exception as e:
        return JSONResponse({"error": str(e)})

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8000)