# app_server.py ✅ Final version (Police Station Complaint)

import os
from flask import Flask, render_template, request, send_file, redirect, url_for, flash
from io import BytesIO
from docx import Document
from datetime import datetime

# ✅ Gemini official SDK
import google.generativeai as genai

# -----------------------------------------------------------
# ✅ ADD YOUR GEMINI API KEY HERE
# -----------------------------------------------------------
GEMINI_API_KEY = "AIzaSyCo3AF8OzpCwbdh3Zirhp1DhlCUJwWApeQ"
# -----------------------------------------------------------

if not GEMINI_API_KEY:
    raise ValueError("❌ GEMINI_API_KEY is empty. Paste your API key inside app_server.py")

genai.configure(api_key=GEMINI_API_KEY)
model = genai.GenerativeModel("gemini-2.0-flash")

app = Flask(__name__)
app.secret_key = "legal-draft-secret-key"

# ---------------------------------------------------------------------
# Build Prompt for drafting Indian Police Station complaint
# ---------------------------------------------------------------------
def build_prompt(data: dict) -> str:
    date_str = data.get("incident_date") or datetime.today().strftime("%d %B %Y")

    prompt = f"""
DRAFT POLICE COMPLAINT (INDIA)

You are a professional legal drafting assistant. Prepare a formal written complaint suitable to submit to a Police Station in India.

Use **only** the user's data — do not invent facts. Follow Indian police complaint format:
- Heading: "To The Officer In-Charge, [Police Station Name]"
- Parties: Complainant / Accused
- Numbered facts of incident
- Prayer / Relief section
- Annexure / Evidence list
- Verification (Affidavit style)

COMPLAINT DETAILS:
Complainant name: {data.get('petitioner_name','')}
Complainant address: {data.get('petitioner_address','')}
Accused / Respondent: {data.get('respondent_name','')}
Accused address: {data.get('respondent_address','')}
Police Station: {data.get('place','')}
Date of incident: {date_str}
Facts in detail: {data.get('facts','')}
Relief sought: {data.get('relief','')}
Sections mentioned (if any): {data.get('sections','')}
Evidence / documents / witnesses: {data.get('evidence','')}

TASKS:
1. Generate a legally formatted Indian Police Station complaint.
2. Include heading, cause title, numbered facts, prayer, annexure list, and verification.
3. Suggest IPC/CrPC sections **only if relevant**. Do not give legal advice.
4. Output should be plain formatted text — NO MARKDOWN.

End with this disclaimer:
"This is a draft. Consult a practising advocate before filing."

START DRAFT BELOW:
"""
    return prompt

# -----------------------------------------------------------
# Gemini request function
# -----------------------------------------------------------
def call_gemini_generate(prompt_text: str) -> str:
    try:
        response = model.generate_content(prompt_text)
        return response.text
    except Exception as e:
        return f"[ERROR] Gemini request failed: {e}"

# -----------------------------------------------------------
# Generate .docx file from text
# -----------------------------------------------------------
def create_docx_from_text(title: str, body_text: str) -> BytesIO:
    doc = Document()
    doc.add_heading(title, level=1)

    for line in body_text.splitlines():
        if line.strip() == "":
            doc.add_paragraph("")
        else:
            doc.add_paragraph(line)

    file_stream = BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    return file_stream

# -----------------------------------------------------------
# Routes
# -----------------------------------------------------------
@app.route("/")
def index():
    return redirect(url_for("complaint_form"))

@app.route("/petition")
def complaint_form():
    return render_template("petition_form.html")  # <-- you should rename petition_form.html to complaint_form.html

@app.route("/generate", methods=["POST"])
def generate():
    data = {
        "petitioner_name": request.form.get("petitioner_name", ""),
        "petitioner_address": request.form.get("petitioner_address", ""),
        "respondent_name": request.form.get("respondent_name", ""),
        "respondent_address": request.form.get("respondent_address", ""),
        "place": request.form.get("place", ""),
        "incident_date": request.form.get("incident_date", ""),
        "facts": request.form.get("facts", ""),
        "relief": request.form.get("relief", ""),
        "sections": request.form.get("sections", ""),
        "evidence": request.form.get("evidence", "")
    }

    prompt = build_prompt(data)
    generated_text = call_gemini_generate(prompt)

    # Save .docx in memory & temp folder
    filename = f"police_complaint_{datetime.now().strftime('%Y%m%d%H%M%S')}.docx"
    file_stream = create_docx_from_text("POLICE COMPLAINT", generated_text)

    os.makedirs("temp_outputs", exist_ok=True)
    filepath = os.path.join("temp_outputs", filename)

    with open(filepath, "wb") as f:
        f.write(file_stream.read())

    file_stream.seek(0)
    return render_template("petition_result.html",  # you can reuse same template
                           draft_text=generated_text,
                           download_filename=filename)

@app.route("/download/<fname>")
def download_file(fname):
    file_path = os.path.join("temp_outputs", fname)
    if not os.path.exists(file_path):
        flash("File not found.", "danger")
        return redirect(url_for("complaint_form"))

    return send_file(file_path, as_attachment=True, download_name=fname)

if __name__ == "__main__":
    app.run(debug=True, port=9000)
